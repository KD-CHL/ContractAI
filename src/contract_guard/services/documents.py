"""Document ingestion and text extraction.

The service keeps binary parsing separate from the review workflow.  Image OCR
is deliberately represented by a small protocol so an on-premise model, a
hosted vision model, or a test double can be supplied without changing the API.
"""

from __future__ import annotations

import hashlib
import inspect
import zipfile
from collections.abc import Awaitable, Mapping
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any, Literal, Protocol, runtime_checkable


class DocumentError(ValueError):
    """Base class for errors safe to expose as a client input error."""


class DocumentTooLargeError(DocumentError):
    """Raised when an upload exceeds the configured limit."""


class UnsupportedDocumentTypeError(DocumentError):
    """Raised when a file is not one of the supported document formats."""


class DocumentParsingError(DocumentError):
    """Raised when a supported document cannot be decoded."""


class DocumentComplexityError(DocumentError):
    """Raised when a document exceeds a safe parsing complexity budget."""


@runtime_checkable
class VisionOCR(Protocol):
    """Injectable OCR boundary used for image and scanned-PDF extraction."""

    def extract_text(
        self,
        image: bytes,
        *,
        media_type: str,
        filename: str | None = None,
    ) -> str | Awaitable[str]: ...


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    """Normalized result of document extraction."""

    text: str
    fingerprint: str
    filename: str
    media_type: str
    size_bytes: int
    page_count: int | None = None
    metadata: Mapping[str, Any] = field(default_factory=dict)

    @property
    def character_count(self) -> int:
        return len(self.text)


_EXTENSION_MEDIA_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".tif": "image/tiff",
    ".tiff": "image/tiff",
    ".bmp": "image/bmp",
    ".gif": "image/gif",
    ".heic": "image/heic",
    ".heif": "image/heif",
}

_SUPPORTED_IMAGE_TYPES = {
    "image/png",
    "image/jpeg",
    "image/webp",
    "image/tiff",
    "image/bmp",
    "image/gif",
}

_GENERIC_MEDIA_TYPES = {"", "application/octet-stream", "binary/octet-stream"}


def _image_dimensions(content: bytes, media_type: str) -> tuple[int, int] | None:
    """Read common raster dimensions without decoding the image payload."""

    if (
        media_type == "image/png"
        and len(content) >= 24
        and content[8:12] == b"\x00\x00\x00\x0d"
        and content[12:16] == b"IHDR"
    ):
        return int.from_bytes(content[16:20], "big"), int.from_bytes(content[20:24], "big")
    if media_type == "image/gif" and len(content) >= 10:
        return int.from_bytes(content[6:8], "little"), int.from_bytes(content[8:10], "little")
    if media_type == "image/bmp" and len(content) >= 26:
        width = abs(int.from_bytes(content[18:22], "little", signed=True))
        height = abs(int.from_bytes(content[22:26], "little", signed=True))
        return width, height
    if media_type == "image/tiff" and len(content) >= 10:
        byte_order = content[:2]
        endian: Literal["little", "big"] | None = (
            "little" if byte_order == b"II" else "big" if byte_order == b"MM" else None
        )
        if endian and int.from_bytes(content[2:4], endian) == 42:
            directory_offset = int.from_bytes(content[4:8], endian)
            if directory_offset + 2 <= len(content):
                entry_count = min(
                    int.from_bytes(
                        content[directory_offset : directory_offset + 2],
                        endian,
                    ),
                    4096,
                )
                dimensions: dict[int, int] = {}
                for index in range(entry_count):
                    start = directory_offset + 2 + index * 12
                    if start + 12 > len(content):
                        break
                    tag = int.from_bytes(content[start : start + 2], endian)
                    value_type = int.from_bytes(content[start + 2 : start + 4], endian)
                    count = int.from_bytes(content[start + 4 : start + 8], endian)
                    if tag not in {256, 257} or count != 1 or value_type not in {3, 4}:
                        continue
                    size = 2 if value_type == 3 else 4
                    dimensions[tag] = int.from_bytes(content[start + 8 : start + 8 + size], endian)
                if 256 in dimensions and 257 in dimensions:
                    return dimensions[256], dimensions[257]
    if media_type == "image/webp" and len(content) >= 30:
        chunk = content[12:16]
        if chunk == b"VP8X":
            width = 1 + int.from_bytes(content[24:27], "little")
            height = 1 + int.from_bytes(content[27:30], "little")
            return width, height
        if chunk == b"VP8 " and content[23:26] == b"\x9d\x01\x2a":
            width = int.from_bytes(content[26:28], "little") & 0x3FFF
            height = int.from_bytes(content[28:30], "little") & 0x3FFF
            return width, height
        if chunk == b"VP8L" and content[20] == 0x2F:
            bits = int.from_bytes(content[21:25], "little")
            return (bits & 0x3FFF) + 1, ((bits >> 14) & 0x3FFF) + 1
    if media_type == "image/jpeg":
        offset = 2
        while offset + 9 <= len(content):
            if content[offset] != 0xFF:
                offset += 1
                continue
            while offset < len(content) and content[offset] == 0xFF:
                offset += 1
            if offset >= len(content):
                break
            marker = content[offset]
            offset += 1
            if marker in {0xD8, 0xD9}:
                continue
            if offset + 2 > len(content):
                break
            segment_size = int.from_bytes(content[offset : offset + 2], "big")
            if segment_size < 2 or offset + segment_size > len(content):
                break
            if (
                marker
                in {
                    0xC0,
                    0xC1,
                    0xC2,
                    0xC3,
                    0xC5,
                    0xC6,
                    0xC7,
                    0xC9,
                    0xCA,
                    0xCB,
                    0xCD,
                    0xCE,
                    0xCF,
                }
                and segment_size >= 7
            ):
                height = int.from_bytes(content[offset + 3 : offset + 5], "big")
                width = int.from_bytes(content[offset + 5 : offset + 7], "big")
                return width, height
            offset += segment_size
    return None


def document_fingerprint(content: bytes) -> str:
    """Return the stable SHA-256 identity of the original document bytes."""

    return hashlib.sha256(content).hexdigest()


def text_fingerprint(text: str) -> str:
    """Return a fingerprint for text submitted without a file."""

    return document_fingerprint(text.encode("utf-8"))


def _sniff_media_type(content: bytes) -> str | None:
    if content.startswith(b"%PDF-"):
        return "application/pdf"
    if content.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if content.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    if content.startswith((b"II*\x00", b"MM\x00*")):
        return "image/tiff"
    if content.startswith(b"BM"):
        return "image/bmp"
    if content.startswith((b"GIF87a", b"GIF89a")):
        return "image/gif"
    if len(content) >= 12 and content[:4] == b"RIFF" and content[8:12] == b"WEBP":
        return "image/webp"
    return None


def _media_type_for(content: bytes, filename: str | None, declared_media_type: str | None) -> str:
    sniffed = _sniff_media_type(content)
    declared = (declared_media_type or "").split(";", 1)[0].strip().lower()
    suffix = Path(filename or "").suffix.lower()
    from_extension = _EXTENSION_MEDIA_TYPES.get(suffix)

    # Binary signatures take precedence over client-supplied headers. DOCX is a
    # ZIP container and is identified by its extension/content type below.
    if sniffed:
        return sniffed
    if declared not in _GENERIC_MEDIA_TYPES:
        return declared
    if from_extension:
        return from_extension
    return declared or "application/octet-stream"


def _decode_plain_text(content: bytes) -> str:
    if b"\x00" in content[:4096] and not content.startswith((b"\xff\xfe", b"\xfe\xff")):
        raise DocumentParsingError("the text document appears to contain binary data")
    for encoding in ("utf-8-sig", "utf-16", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParsingError("the text document encoding is not supported")


def extract_pdf_text(
    content: bytes,
    *,
    max_pages: int = 200,
) -> tuple[str, int, dict[str, Any]]:
    """Extract searchable text from a PDF using PyMuPDF."""

    try:
        import fitz  # type: ignore[import-untyped]
    except ImportError as exc:  # pragma: no cover - dependency error path
        raise RuntimeError("PDF support requires the PyMuPDF package") from exc

    try:
        with fitz.open(stream=content, filetype="pdf") as pdf:
            if len(pdf) > max_pages:
                raise DocumentComplexityError(f"PDF 页数超过 {max_pages} 页限制")
            page_text = [page.get_text("text").strip() for page in pdf]
            metadata = {
                str(key): value
                for key, value in (pdf.metadata or {}).items()
                if value not in (None, "")
            }
            return "\n\n".join(part for part in page_text if part), len(pdf), metadata
    except DocumentError:
        raise
    except Exception as exc:
        raise DocumentParsingError("the PDF could not be parsed") from exc


def extract_docx_text(
    content: bytes,
    *,
    max_uncompressed_bytes: int = 100 * 1024 * 1024,
    max_entries: int = 10_000,
) -> tuple[str, dict[str, Any]]:
    """Extract paragraphs and table cells from a DOCX document."""

    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency error path
        raise RuntimeError("DOCX support requires the python-docx package") from exc

    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            entries = archive.infolist()
            if len(entries) > max_entries:
                raise DocumentComplexityError("DOCX 内部文件数量超过安全限制")
            expanded_size = sum(item.file_size for item in entries)
            if expanded_size > max_uncompressed_bytes:
                raise DocumentComplexityError(
                    f"DOCX 解压后大小超过 {max_uncompressed_bytes} 字节限制"
                )
        document = Document(BytesIO(content))
        blocks: list[str] = []
        blocks.extend(p.text.strip() for p in document.paragraphs if p.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append("\t".join(cells))

        properties = document.core_properties
        metadata = {
            key: value
            for key, value in {
                "title": properties.title,
                "subject": properties.subject,
                "author": properties.author,
                "keywords": properties.keywords,
            }.items()
            if value
        }
        return "\n".join(blocks), metadata
    except DocumentError:
        raise
    except Exception as exc:
        raise DocumentParsingError("the DOCX document could not be parsed") from exc


class DocumentService:
    """Validate an uploaded document and turn it into reviewable text."""

    def __init__(
        self,
        *,
        vision_ocr: VisionOCR | None = None,
        max_bytes: int = 20 * 1024 * 1024,
        max_pages: int = 200,
        max_characters: int = 1_000_000,
        max_docx_uncompressed_bytes: int = 100 * 1024 * 1024,
        max_ocr_pixels: int = 20_000_000,
    ) -> None:
        if (
            min(
                max_bytes,
                max_pages,
                max_characters,
                max_docx_uncompressed_bytes,
                max_ocr_pixels,
            )
            <= 0
        ):
            raise ValueError("document limits must be greater than zero")
        self.vision_ocr = vision_ocr
        self.max_bytes = max_bytes
        self.max_pages = max_pages
        self.max_characters = max_characters
        self.max_docx_uncompressed_bytes = max_docx_uncompressed_bytes
        self.max_ocr_pixels = max_ocr_pixels

    async def parse(
        self,
        content: bytes,
        *,
        filename: str | None = None,
        media_type: str | None = None,
    ) -> ParsedDocument:
        if not content:
            raise DocumentParsingError("the uploaded document is empty")
        if len(content) > self.max_bytes:
            raise DocumentTooLargeError(
                f"the uploaded document exceeds the {self.max_bytes}-byte limit"
            )

        safe_name = Path(filename or "document").name or "document"
        resolved_type = _media_type_for(content, safe_name, media_type)
        metadata: dict[str, Any] = {}
        page_count: int | None = None

        if resolved_type == "application/pdf":
            text, page_count, metadata = extract_pdf_text(content, max_pages=self.max_pages)
            if not text and self.vision_ocr is not None:
                text = await self._ocr_pdf(content, safe_name)
        elif resolved_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ) or safe_name.lower().endswith(".docx"):
            text, metadata = extract_docx_text(
                content,
                max_uncompressed_bytes=self.max_docx_uncompressed_bytes,
            )
            resolved_type = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        elif resolved_type in {"text/plain", "text/markdown"}:
            text = _decode_plain_text(content)
        elif resolved_type in _SUPPORTED_IMAGE_TYPES or resolved_type.startswith("image/"):
            text = await self._ocr_image(content, media_type=resolved_type, filename=safe_name)
            page_count = 1
        else:
            raise UnsupportedDocumentTypeError(f"unsupported document type: {resolved_type}")

        text = text.strip()
        if not text:
            raise DocumentParsingError("no reviewable text could be extracted")
        if len(text) > self.max_characters:
            raise DocumentComplexityError(f"可审阅文字超过 {self.max_characters} 字符限制")

        return ParsedDocument(
            text=text,
            fingerprint=document_fingerprint(content),
            filename=safe_name,
            media_type=resolved_type,
            size_bytes=len(content),
            page_count=page_count,
            metadata=metadata,
        )

    async def _ocr_image(self, content: bytes, *, media_type: str, filename: str | None) -> str:
        if self.vision_ocr is None:
            raise UnsupportedDocumentTypeError(
                "image documents require a configured VisionOCR implementation"
            )
        dimensions = _image_dimensions(content, media_type)
        if dimensions is not None:
            width, height = dimensions
            if width <= 0 or height <= 0:
                raise DocumentParsingError("the image dimensions are invalid")
            if width * height > self.max_ocr_pixels:
                raise DocumentComplexityError(f"图片像素超过 {self.max_ocr_pixels} 安全限制")
        try:
            result = self.vision_ocr.extract_text(content, media_type=media_type, filename=filename)
            if inspect.isawaitable(result):
                result = await result
        except ValueError as exc:
            raise DocumentParsingError(str(exc)) from exc
        if not isinstance(result, str):
            raise DocumentParsingError("VisionOCR returned a non-text result")
        return result

    async def _ocr_pdf(self, content: bytes, filename: str) -> str:
        try:
            import fitz
        except ImportError as exc:  # pragma: no cover - dependency error path
            raise RuntimeError("PDF support requires the PyMuPDF package") from exc

        extracted: list[str] = []
        try:
            with fitz.open(stream=content, filetype="pdf") as pdf:
                if len(pdf) > self.max_pages:
                    raise DocumentComplexityError(f"PDF 页数超过 {self.max_pages} 页限制")
                for page_number, page in enumerate(pdf, start=1):
                    rendered_width = int(page.rect.width * 2)
                    rendered_height = int(page.rect.height * 2)
                    if rendered_width * rendered_height > self.max_ocr_pixels:
                        raise DocumentComplexityError(
                            f"PDF 第 {page_number} 页渲染尺寸超过安全限制"
                        )
                    pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    value = await self._ocr_image(
                        pixmap.tobytes("png"),
                        media_type="image/png",
                        filename=f"{filename}#page={page_number}",
                    )
                    if value.strip():
                        extracted.append(value.strip())
        except DocumentError:
            raise
        except Exception as exc:
            raise DocumentParsingError("the scanned PDF could not be OCR processed") from exc
        return "\n\n".join(extracted)


# Backwards-friendly name for callers that prefer a parser-oriented abstraction.
DocumentParser = DocumentService
