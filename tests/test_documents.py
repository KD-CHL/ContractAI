from io import BytesIO

import fitz
import pytest
from docx import Document

from contract_guard.services.documents import (
    DocumentComplexityError,
    DocumentService,
    DocumentTooLargeError,
    UnsupportedDocumentTypeError,
    document_fingerprint,
)


class FakeOCR:
    def __init__(self) -> None:
        self.calls: list[tuple[str, str | None]] = []

    async def extract_text(
        self, image: bytes, *, media_type: str, filename: str | None = None
    ) -> str:
        assert image
        self.calls.append((media_type, filename))
        return "OCR extracted contract text"


@pytest.mark.parametrize(
    ("filename", "media_type"),
    [("contract.txt", "text/plain"), ("contract.md", "text/markdown")],
)
async def test_parse_plain_text_and_markdown(filename: str, media_type: str) -> None:
    content = "甲方应于三十日内付款。".encode()

    parsed = await DocumentService().parse(content, filename=filename)

    assert parsed.text == "甲方应于三十日内付款。"
    assert parsed.media_type == media_type
    assert parsed.fingerprint == document_fingerprint(content)


async def test_parse_pdf_with_pymupdf() -> None:
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Payment is due within 30 days.")
    content = pdf.tobytes()
    pdf.close()

    parsed = await DocumentService().parse(content, filename="terms.pdf")

    assert "Payment is due within 30 days" in parsed.text
    assert parsed.page_count == 1
    assert parsed.media_type == "application/pdf"


async def test_parse_docx_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("Service terms")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Payment"
    table.cell(0, 1).text = "30 days"
    output = BytesIO()
    document.save(output)

    parsed = await DocumentService().parse(output.getvalue(), filename="terms.docx")

    assert "Service terms" in parsed.text
    assert "Payment\t30 days" in parsed.text


async def test_image_uses_injected_vision_ocr() -> None:
    ocr = FakeOCR()
    content = b"\x89PNG\r\n\x1a\nnot-a-real-image-needed-by-the-fake"

    parsed = await DocumentService(vision_ocr=ocr).parse(content, filename="scan.png")

    assert parsed.text == "OCR extracted contract text"
    assert parsed.page_count == 1
    assert ocr.calls == [("image/png", "scan.png")]


async def test_image_without_ocr_is_rejected() -> None:
    with pytest.raises(UnsupportedDocumentTypeError, match="VisionOCR"):
        await DocumentService().parse(b"\xff\xd8\xffimage", filename="scan.jpg")


async def test_image_pixel_limit_is_enforced_before_ocr() -> None:
    width = 5000
    height = 5000
    content = (
        b"\x89PNG\r\n\x1a\n"
        + b"\x00\x00\x00\x0dIHDR"
        + width.to_bytes(4, "big")
        + height.to_bytes(4, "big")
    )

    with pytest.raises(DocumentComplexityError, match="20000000"):
        await DocumentService(vision_ocr=FakeOCR()).parse(content, filename="large.png")


async def test_upload_limit_is_enforced() -> None:
    with pytest.raises(DocumentTooLargeError):
        await DocumentService(max_bytes=3).parse(b"four", filename="contract.txt")


async def test_character_limit_is_enforced() -> None:
    with pytest.raises(DocumentComplexityError, match="10"):
        await DocumentService(max_characters=10).parse(
            "这是一份明显超过十个字符的合同文本".encode(),
            filename="contract.txt",
        )


async def test_pdf_page_limit_is_enforced() -> None:
    pdf = fitz.open()
    pdf.new_page()
    pdf.new_page()
    content = pdf.tobytes()
    pdf.close()

    with pytest.raises(DocumentComplexityError, match="1"):
        await DocumentService(max_pages=1).parse(content, filename="terms.pdf")


async def test_docx_expanded_size_limit_is_enforced() -> None:
    document = Document()
    document.add_paragraph("Service terms")
    output = BytesIO()
    document.save(output)

    with pytest.raises(DocumentComplexityError, match="1024"):
        await DocumentService(max_docx_uncompressed_bytes=1024).parse(
            output.getvalue(),
            filename="terms.docx",
        )
